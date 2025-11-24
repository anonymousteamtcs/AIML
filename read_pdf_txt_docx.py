from pathlib import Path
from typing import Tuple
import io

# Optional libraries; ensure installed with pip
from docx import Document
from PyPDF2 import PdfReader

def read_file_text(path: str, *, return_bytes_on_failure: bool = False) -> Tuple[str, bool]:
    """
    Read text from many common file types and return the combined text and a success flag.
    The function also prints the content.

    Parameters:
    - path: filesystem path to the file
    - return_bytes_on_failure: if True and text decoding fails, returns the raw bytes decoded with latin-1

    Returns:
    - (content, success) where content is a string and success is True when read as text normally.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = p.suffix.lower()

    def _print_and_return(s: str, ok: bool = True):
        print(s)
        return s, ok

    # Plain text-like files
    if suffix in {".txt", ".log", ".md", ".csv", ".json", ".py", ".ini", ".cfg", ".yaml", ".yml"}:
        try:
            text = p.read_text(encoding="utf-8")
            return _print_and_return(text)
        except UnicodeDecodeError:
            # fallback to common encodings
            for enc in ("latin-1", "cp1252"):
                try:
                    text = p.read_text(encoding=enc)
                    return _print_and_return(text)
                except Exception:
                    continue
            if return_bytes_on_failure:
                raw = p.read_bytes().decode("latin-1", errors="replace")
                return _print_and_return(raw, False)
            raise

    # DOCX
    if suffix == ".docx":
        try:
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            # Also capture text inside tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            paragraphs.append(cell.text)
            text = "\n".join(paragraphs)
            return _print_and_return(text)
        except Exception as e:
            raise RuntimeError(f"Failed to read docx: {e}")

    # PDF
    if suffix == ".pdf":
        try:
            reader = PdfReader(path)
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    # continue on page error
                    pages.append("")
            text = "\n".join(pages)
            return _print_and_return(text)
        except Exception as e:
            raise RuntimeError(f"Failed to read pdf: {e}")

    # RTF - quick best-effort: treat as text, strip common RTF tokens
    if suffix == ".rtf":
        try:
            raw = p.read_text(encoding="utf-8", errors="replace")
            # very light RTF cleanup
            import re
            cleaned = re.sub(r"\\[a-zA-Z]+\d* ?", "", raw)        # remove control words
            cleaned = re.sub(r"\{\\.*?\}", "", cleaned)          # remove groupings (naive)
            cleaned = cleaned.replace("{", "").replace("}", "")
            return _print_and_return(cleaned)
        except Exception:
            raise

    # Common office or binary formats not directly supported: try text fallback
    try:
        # Try reading as UTF-8 text first
        text = p.read_text(encoding="utf-8")
        return _print_and_return(text)
    except Exception:
        # If binary, try decoding bytes with a tolerant codec
        try:
            raw = p.read_bytes()
            decoded = raw.decode("utf-8", errors="replace")
            return _print_and_return(decoded, False)
        except Exception as e:
            if return_bytes_on_failure:
                # force latin-1 decode (one-to-one byte->char)
                forced = p.read_bytes().decode("latin-1", errors="replace")
                return _print_and_return(forced, False)
            raise RuntimeError(f"Unable to read file {path}: {e}")
        

# if __name__ == "__main__":
#     path = "sample-report.pdf"
#     content, success = read_file_text(path)
#     if success:
#         print("File read successfully:")
#     else:
#         print("File read with issues:")
#     print(content)
